"""
CV Parser Service - Extract structured data from CV files using AI.
"""
import io
import json
import logging
import re
from datetime import datetime

from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from app.domains.cv_screening.repositories.screening_repository import ScreeningRepository
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from lia_models.candidate import Candidate
from app.schemas.cv_parser import DuplicateCheck, Education, Experience, ParsedCV
from app.domains.ai.services.llm import LLMService
from app.shared.robustness.document_scanner import scan_document

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
SUPPORTED_FORMATS = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
    'text/plain': 'txt',
}


class CVParserService:
    """
    Service for parsing CVs and extracting structured candidate information.
    Uses Claude AI for intelligent extraction.
    """
    
    def __init__(self):
        self.llm_service = LLMService()
    
    async def parse_cv(self, file: UploadFile) -> ParsedCV:
        """
        Parse an uploaded CV file and extract structured data.
        
        Args:
            file: Uploaded file (PDF, DOCX, DOC, or TXT)
        
        Returns:
            ParsedCV with extracted information
        
        Raises:
            HTTPException if file is invalid or parsing fails
        """
        content = await file.read()
        file_size = len(content)
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="Empty file")
        
        content_type = file.content_type or ''
        file_extension = self._get_file_extension(file.filename or '')
        
        file_type = SUPPORTED_FORMATS.get(content_type)
        if not file_type and file_extension in ['pdf', 'docx', 'doc', 'txt']:
            file_type = file_extension
        
        if not file_type:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file format. Supported: PDF, DOCX, DOC, TXT"
            )
        
        try:
            if file_type == 'pdf':
                text = await self.extract_text_from_pdf(content)
            elif file_type in ['docx', 'doc']:
                text = await self.extract_text_from_docx(content)
            elif file_type == 'txt':
                text = content.decode('utf-8', errors='ignore')
            else:
                raise HTTPException(status_code=400, detail="Unsupported file format")
            
            if len(text.strip()) < 50:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract enough text from file. The file may be empty, corrupted, or image-based."
                )
            
            # Scan document for injection payloads before AI processing
            scan_result = scan_document(text)
            if scan_result.threats:
                logger.warning("CV document scan threats: %s", scan_result.threats)
            text = scan_result.sanitized_text
            parsed_cv = await self.extract_with_ai(text)
            
            parsed_cv.file_name = file.filename
            parsed_cv.file_type = file_type
            parsed_cv.file_size_bytes = file_size
            parsed_cv.raw_text = text
            
            return parsed_cv
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"CV parsing failed: {e}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail="Internal server error"
            )
    
    async def extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to read PDF: {str(e)}"
            )
    
    async def extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to read DOCX: {str(e)}"
            )
    
    async def extract_with_ai(self, text: str) -> ParsedCV:
        """
        Use Claude AI to extract structured information from CV text.
        Optimized for Brazilian CVs.
        """
        prompt = f"""Você é um especialista em extração de dados de currículos brasileiros. Analise o texto do CV abaixo e extraia TODAS as informações estruturadas disponíveis.

INSTRUÇÕES CRÍTICAS PARA EXTRAÇÃO:

1. **NOME COMPLETO** (OBRIGATÓRIO):
   - O nome geralmente está nas PRIMEIRAS LINHAS do CV, em destaque
   - Procure por nomes próprios brasileiros (ex: "João Silva", "Maria Santos", "Carlos Eduardo")
   - NUNCA retorne "Unknown", "Nome não identificado" ou null se houver qualquer nome no texto
   - Se encontrar múltiplos nomes, use o primeiro que aparece no topo/cabeçalho

2. **TELEFONE** (formatos brasileiros):
   - Formatos válidos: +55 11 99999-9999, (11) 99999-9999, 11 99999-9999, 55 11 999999999
   - Celulares brasileiros têm 9 dígitos após o DDD (começam com 9)
   - DDDs comuns: 11 (SP), 21 (RJ), 31 (BH), 41 (Curitiba), 51 (POA), etc.
   - Normalize para formato: +55 XX XXXXX-XXXX

3. **LOCALIZAÇÃO**:
   - Extraia cidade, estado e país quando disponível
   - Formatos comuns: "São Paulo - SP", "Rio de Janeiro, RJ", "Belo Horizonte/MG"
   - Retorne no formato: "Cidade, Estado, Brasil" ou "Cidade, Estado"

4. **CARGO ATUAL** (current_title):
   - Identifique o cargo da experiência mais recente
   - Palavras que indicam posição atual: "atual", "presente", "current", "present", "até o momento"
   - Se não houver indicação de "atual", use a primeira experiência listada

5. **HABILIDADES TÉCNICAS** (skills):
   - Extraia APENAS tecnologias, linguagens, frameworks e ferramentas técnicas
   - Inclua: linguagens de programação (Python, Java, JavaScript, C#, etc.)
   - Inclua: frameworks (React, Angular, Django, Spring, .NET, etc.)
   - Inclua: bancos de dados (PostgreSQL, MySQL, MongoDB, Redis, etc.)
   - Inclua: ferramentas (Docker, Kubernetes, Git, AWS, Azure, etc.)
   - Inclua: metodologias técnicas (Scrum, Agile, DevOps, CI/CD, etc.)
   - NÃO inclua soft skills aqui - coloque em soft_skills

6. **SOFT SKILLS** (soft_skills):
   - Extraia habilidades interpessoais e comportamentais separadamente
   - Inclua: comunicação, liderança, trabalho em equipe, negociação
   - Inclua: resolução de problemas, pensamento crítico, criatividade
   - Inclua: organização, gestão de tempo, adaptabilidade, proatividade
   - Inclua: inteligência emocional, empatia, colaboração

7. **NÍVEL DE SENIORIDADE** (seniority_level):
   - Analise o título atual e anos de experiência para determinar:
     - "junior" ou "júnior": 0-2 anos ou títulos com "Junior/Jr"
     - "pleno": 2-5 anos ou sem prefixo de senioridade
     - "senior" ou "sênior": 5-8 anos ou títulos com "Senior/Sr"
     - "lead" ou "tech lead": liderança técnica
     - "manager": gestão de pessoas
     - "director": direção de área
     - "c-level": C-suite (CEO, CTO, CFO, etc.)

8. **DATA DE NASCIMENTO** (date_of_birth):
   - Procure por: "Nascimento:", "Data de Nascimento:", "Idade:", formatos DD/MM/AAAA
   - Retorne no formato: "DD/MM/AAAA" ou null se não encontrado

9. **INTERESSES** (interests):
   - Procure seções como: "Interesses", "Hobbies", "Atividades", "Sobre mim"
   - Extraia interesses pessoais e profissionais mencionados

Retorne um JSON com esta estrutura EXATA:

{{
    "full_name": "Nome completo do candidato (NUNCA null ou Unknown)",
    "email": "email@exemplo.com ou null",
    "phone": "+55 XX XXXXX-XXXX ou null",
    "linkedin": "URL completa do LinkedIn ou null",
    "github": "URL completa do GitHub ou null",
    "portfolio": "URL do portfólio/website ou null",
    "location": "Cidade, Estado, País ou null",
    "current_title": "Cargo atual do candidato ou null",
    "summary": "Resumo profissional/objetivo ou null",
    "experiences": [
        {{
            "company": "Nome da empresa",
            "title": "Cargo/Título",
            "start_date": "Mês Ano (ex: 'Jan 2020' ou 'Janeiro 2020')",
            "end_date": "Mês Ano ou 'Atual'",
            "is_current": true se for o cargo atual,
            "description": "Descrição das atividades",
            "location": "Local de trabalho ou null"
        }}
    ],
    "education": [
        {{
            "institution": "Nome da instituição",
            "degree": "Tipo de graduação (Bacharelado, Mestrado, MBA, Técnico, etc.)",
            "field_of_study": "Área/Curso",
            "start_date": "Ano início ou null",
            "end_date": "Ano fim ou null",
            "is_completed": true/false,
            "description": "Detalhes adicionais ou null"
        }}
    ],
    "skills": ["Python", "React", "AWS", "Docker", "...APENAS skills técnicas"],
    "soft_skills": ["Liderança", "Comunicação", "Trabalho em equipe", "..."],
    "languages": ["Português (Nativo)", "Inglês (Fluente)", "..."],
    "certifications": ["Certificação 1", "Certificação 2"],
    "seniority_level": "junior|pleno|senior|lead|manager|director|c-level ou null",
    "date_of_birth": "DD/MM/AAAA ou null",
    "interests": ["Interesse 1", "Interesse 2", "..."],
    "confidence_score": 0.0 a 1.0 (baseado na completude do CV),
    "extraction_notes": ["Notas sobre campos que podem precisar revisão"]
}}

REGRAS FINAIS:
- Experiências devem estar ordenadas da mais recente para a mais antiga
- is_current=true APENAS para o cargo atual (com "atual", "presente", "present")
- Se o CV estiver incompleto, extraia o máximo possível e ajuste confidence_score
- Separe CLARAMENTE skills técnicas de soft skills
- Não invente informações - use null ou lista vazia para campos não encontrados

TEXTO DO CV:
{text[:15000]}

Retorne APENAS o objeto JSON, sem texto adicional ou markdown."""

        try:
            response = await self.llm_service.generate(
                prompt=prompt,
            )
            
            json_str = response.strip()
            if json_str.startswith('```json'):
                json_str = json_str[7:]
            if json_str.startswith('```'):
                json_str = json_str[3:]
            if json_str.endswith('```'):
                json_str = json_str[:-3]
            
            data = json.loads(json_str.strip())
            
            experiences = [
                Experience(**exp) for exp in data.get('experiences', [])
            ]
            education = [
                Education(**edu) for edu in data.get('education', [])
            ]
            
            current_title = data.get('current_title')
            if not current_title and experiences:
                current_title = experiences[0].title
            
            parsed_cv = ParsedCV(
                full_name=data.get('full_name', 'Unknown'),
                email=data.get('email'),
                phone=data.get('phone'),
                linkedin=data.get('linkedin'),
                github=data.get('github'),
                portfolio=data.get('portfolio'),
                location=data.get('location'),
                current_title=current_title,
                summary=data.get('summary'),
                experiences=experiences,
                education=education,
                skills=data.get('skills', []),
                languages=data.get('languages', []),
                certifications=data.get('certifications', []),
                soft_skills=data.get('soft_skills', []),
                seniority_level=data.get('seniority_level'),
                date_of_birth=data.get('date_of_birth'),
                interests=data.get('interests', []),
                raw_text=text,
                file_name=None,
                file_type=None,
                file_size_bytes=None,
                file_url=None,
                confidence_score=float(data.get('confidence_score', 0.7)),
                extraction_notes=data.get('extraction_notes', []),
                parsed_at=datetime.utcnow()
            )
            
            return parsed_cv
            
        except json.JSONDecodeError as e:
            logger.error(f"AI response parsing failed: {e}")
            return await self._fallback_extraction(text)
        except Exception as e:
            logger.error(f"AI extraction failed: {e}", exc_info=True)
            return await self._fallback_extraction(text)
    
    async def _fallback_extraction(self, text: str) -> ParsedCV:
        """
        Fallback extraction using regex patterns when AI fails.
        Optimized for Brazilian CVs.
        """
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        
        br_phone_patterns = [
            r'\+55\s*\(?\d{2}\)?\s*\d{4,5}[-\s]?\d{4}',
            r'\(?\d{2}\)?\s*9?\d{4}[-\s]?\d{4}',
            r'55\s*\d{2}\s*9?\d{4,5}[-\s]?\d{4}',
            r'\d{2}\s*9\d{4}[-\s]?\d{4}',
        ]
        
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        github_pattern = r'github\.com/[\w\-]+'
        
        location_patterns = [
            r'([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-ZÀ-Ú][a-zà-ú]+)*)\s*[-–/,]\s*([A-Z]{2})',
            r'([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-ZÀ-Ú][a-zà-ú]+)*)\s*,\s*([A-ZÀ-Ú][a-zà-ú]+)',
        ]
        
        common_skills = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C#', 'C\\+\\+', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin',
            'React', 'Angular', 'Vue', 'Node\\.js', 'Django', 'Flask', 'FastAPI', 'Spring', '\\.NET', 'Laravel',
            'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git', 'CI/CD', 'Jenkins', 'Terraform',
            'Scrum', 'Agile', 'Kanban', 'DevOps', 'Linux', 'Windows Server',
            'HTML', 'CSS', 'SASS', 'REST', 'GraphQL', 'API', 'Microservices',
            'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Pandas', 'NumPy',
            'Excel', 'Power BI', 'Tableau', 'SAP', 'Salesforce', 'Jira', 'Confluence'
        ]
        
        emails = re.findall(email_pattern, text)
        
        phones = []
        for pattern in br_phone_patterns:
            found = re.findall(pattern, text)
            phones.extend(found)
        
        phone = None
        if phones:
            raw_phone = phones[0]
            digits = re.sub(r'\D', '', raw_phone)
            if len(digits) >= 10:
                if not digits.startswith('55'):
                    digits = '55' + digits
                if len(digits) == 12:
                    digits = digits[:4] + '9' + digits[4:]
                phone = f"+{digits[:2]} {digits[2:4]} {digits[4:9]}-{digits[9:]}"
        
        linkedins = re.findall(linkedin_pattern, text, re.IGNORECASE)
        githubs = re.findall(github_pattern, text, re.IGNORECASE)
        
        location = None
        for pattern in location_patterns:
            matches = re.findall(pattern, text)
            if matches:
                city, state = matches[0]
                location = f"{city}, {state}, Brasil"
                break
        
        skills = []
        text_upper = text
        for skill in common_skills:
            if re.search(rf'\b{skill}\b', text_upper, re.IGNORECASE):
                clean_skill = skill.replace('\\', '')
                if clean_skill not in skills:
                    skills.append(clean_skill)
        
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        name = None
        for line in lines[:10]:
            if len(line) < 5 or len(line) > 60:
                continue
            if '@' in line or 'http' in line.lower():
                continue
            if re.search(r'\d{5,}', line):
                continue
            words = line.split()
            if len(words) >= 2 and all(w[0].isupper() for w in words if w.isalpha()):
                name = line
                break
        
        if not name:
            name = lines[0] if lines else 'Não identificado'
            if len(name) > 60 or '@' in name:
                name = 'Não identificado'
        
        # Try to extract current title from common job title patterns
        current_title = None
        job_title_patterns = [
            r'(?:cargo|posição|função|título)\s*[:\-]?\s*([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-Za-zÀ-ú]+){0,4})',
            r'(?:Desenvolvedor|Developer|Analista|Gerente|Coordenador|Diretor|Especialista|Consultor|Engenheiro|Arquiteto|Designer|Product|Tech Lead|Manager|Senior|Sênior|Pleno|Júnior|Junior)(?:\s+[A-Za-zÀ-ú]+){0,3}',
            r'([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-Za-zÀ-ú]+){0,3})\s*[-–|]\s*(?:Atual|Present|Current)',
        ]
        
        for pattern in job_title_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                title_candidate = matches[0] if isinstance(matches[0], str) else matches[0][0] if matches[0] else None
                if title_candidate and len(title_candidate) > 3 and len(title_candidate) < 60:
                    current_title = title_candidate.strip()
                    break
        
        extraction_notes = [
            "Extração por IA falhou - usando extração por padrões regex",
            "Por favor, revise e complete todos os campos manualmente"
        ]
        if not skills:
            extraction_notes.append("Nenhuma habilidade técnica identificada automaticamente")
        if not location:
            extraction_notes.append("Localização não identificada")
        if not current_title:
            extraction_notes.append("Cargo atual não identificado automaticamente")
        
        return ParsedCV(
            full_name=name,
            email=emails[0] if emails else None,
            phone=phone,
            linkedin=f"https://{linkedins[0]}" if linkedins else None,
            github=f"https://{githubs[0]}" if githubs else None,
            portfolio=None,
            location=location,
            current_title=current_title,
            summary=None,
            experiences=[],
            education=[],
            skills=skills,
            languages=["Português (Nativo)"] if re.search(r'\bportugu[eê]s\b', text, re.IGNORECASE) else [],
            certifications=[],
            soft_skills=[],
            seniority_level=None,
            date_of_birth=None,
            interests=[],
            raw_text=text,
            file_name=None,
            file_type=None,
            file_size_bytes=None,
            file_url=None,
            confidence_score=0.4 if skills else 0.3,
            extraction_notes=extraction_notes,
            parsed_at=datetime.utcnow()
        )
    
    async def check_duplicate(
        self, 
        parsed_cv: ParsedCV, 
        db: AsyncSession
    ) -> DuplicateCheck:
        """
        Check if a candidate with similar information already exists.
        """
        try:
            if parsed_cv.email:
                _cv_email = parsed_cv.email
                existing = await ScreeningRepository(db).find_active_candidate_by_email(_cv_email)
                if existing:
                    candidate_name = getattr(existing, 'name', None)
                    return DuplicateCheck(
                        is_duplicate=True,
                        match_type="email",
                        existing_candidate_id=existing.id if existing.id else None,  # type: ignore[union-attr]
                        existing_candidate_name=str(candidate_name) if candidate_name else None,
                        similarity_score=1.0
                    )
            
            if parsed_cv.full_name and parsed_cv.phone:
                name_normalized = parsed_cv.full_name.lower().strip()
                phone_normalized = re.sub(r'\D', '', parsed_cv.phone)
                
                candidates = await ScreeningRepository(db).list_active_candidates_by_name_lowered(
                    name_normalized
                )
                
                for candidate in candidates:
                    phone_attr = getattr(candidate, 'phone', None)
                    candidate_phone = str(phone_attr) if phone_attr else None
                    if candidate_phone:
                        existing_phone = re.sub(r'\D', '', candidate_phone)
                        if existing_phone and phone_normalized.endswith(existing_phone[-8:]):
                            cand_name = getattr(candidate, 'name', None)
                            return DuplicateCheck(
                                is_duplicate=True,
                                match_type="name_phone",
                                existing_candidate_id=candidate.id if candidate.id else None,  # type: ignore[union-attr]
                                existing_candidate_name=str(cand_name) if cand_name else None,
                                similarity_score=0.9
                            )
            
            if parsed_cv.linkedin:
                linkedin_username = self._extract_linkedin_username(parsed_cv.linkedin)
                if linkedin_username:
                    existing = await ScreeningRepository(db).find_active_candidate_by_linkedin_username(
                        linkedin_username
                    )
                    if existing:
                        exist_name = getattr(existing, 'name', None)
                        return DuplicateCheck(
                            is_duplicate=True,
                            match_type="linkedin",
                            existing_candidate_id=existing.id if existing.id else None,  # type: ignore[union-attr]
                            existing_candidate_name=str(exist_name) if exist_name else None,
                            similarity_score=0.95
                        )
            
            return DuplicateCheck(is_duplicate=False, match_type=None)
            
        except Exception as e:
            logger.error(f"Duplicate check failed: {e}", exc_info=True)
            return DuplicateCheck(is_duplicate=False, match_type=None)
    
    def _extract_linkedin_username(self, url: str) -> str | None:
        """Extract LinkedIn username from URL."""
        match = re.search(r'linkedin\.com/in/([^/?\s]+)', url, re.IGNORECASE)
        return match.group(1) if match else None
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        if '.' in filename:
            return filename.rsplit('.', 1)[-1].lower()
        return ''
    
    def get_supported_formats(self) -> dict:
        """Return supported file formats information."""
        return {
            "formats": [
                {"extension": "pdf", "mime_type": "application/pdf", "name": "PDF Document"},
                {"extension": "docx", "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "name": "Word Document (DOCX)"},
                {"extension": "doc", "mime_type": "application/msword", "name": "Word Document (DOC)"},
                {"extension": "txt", "mime_type": "text/plain", "name": "Plain Text"},
            ],
            "max_file_size_mb": MAX_FILE_SIZE / (1024 * 1024),
            "max_file_size_bytes": MAX_FILE_SIZE
        }


cv_parser_service = CVParserService()

# FastAPI dependency injection factory — returns singleton (holds HTTP clients)
def get_cv_parser_service() -> "CVParserService":
    return cv_parser_service

